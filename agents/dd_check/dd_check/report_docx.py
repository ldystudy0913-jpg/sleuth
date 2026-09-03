"""Build a .docx check report from structured findings. Layout labels come from rubric."""
from __future__ import annotations

from datetime import datetime, timezone
from typing import Any, Dict, List, Mapping


def render_docx_bytes(
    *,
    rubric: Mapping[str, Any],
    score: float,
    summary: str,
    findings: List[Dict[str, Any]],
    sources: List[Dict[str, Any]],
    generated_at: str = "",
) -> bytes:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required: pip install python-docx") from exc

    word = rubric.get("word") if isinstance(rubric.get("word"), dict) else {}
    title = str(word.get("title") or "")
    labels = rubric.get("severity_labels") if isinstance(rubric.get("severity_labels"), dict) else {}
    dim_label = {
        str(d.get("id")): str(d.get("label") or d.get("id"))
        for d in (rubric.get("dimensions") or [])
        if isinstance(d, dict)
    }
    when = generated_at or datetime.now(timezone.utc).strftime(str(word.get("date_format") or "%Y%m%d"))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Calibri")

    doc.add_heading(title, level=0)
    doc.add_paragraph(f"总分：{score}")
    doc.add_paragraph(f"生成日期：{when}")
    if summary:
        doc.add_heading("结论", level=1)
        doc.add_paragraph(summary)

    doc.add_heading("发现问题", level=1)
    if not findings:
        doc.add_paragraph("未列出具体问题。")
    else:
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "严重程度"
        hdr[1].text = "维度"
        hdr[2].text = "位置"
        hdr[3].text = "问题"
        hdr[4].text = "依据"
        for item in findings:
            row = table.add_row().cells
            sev = str(item.get("severity") or "")
            row[0].text = str(labels.get(sev) or sev)
            dim = str(item.get("dimension") or "")
            row[1].text = dim_label.get(dim, dim)
            row[2].text = str(item.get("location") or "")
            row[3].text = str(item.get("issue") or "")
            row[4].text = str(item.get("evidence") or "")

    doc.add_heading("知识来源", level=1)
    if not sources:
        doc.add_paragraph("本次未引用知识库。")
    else:
        for src in sources:
            title_s = str(src.get("title") or src.get("file_name") or "")
            url = str(src.get("url") or "")
            doc.add_paragraph(f"{title_s} {url}".strip())

    buf = __import__("io").BytesIO()
    doc.save(buf)
    return buf.getvalue()
