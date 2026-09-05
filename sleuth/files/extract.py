"""In-process text extraction from decrypted session-file bytes."""
from __future__ import annotations

import base64
import io
from dataclasses import dataclass
from typing import Optional

from ..config import Config, FilesConfig, parse_model_ref


def _files_cfg(config: Optional[Config]) -> FilesConfig:
    if config is None:
        return FilesConfig()
    return getattr(config, "files", None) or FilesConfig()


@dataclass
class Excerpt:
    text: str = ""
    truncated: bool = False
    parser: str = ""
    skipped: str = ""


def _clip(text: str, max_chars: int) -> tuple[str, bool]:
    if max_chars <= 0 or len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


def _ext(filename: str) -> str:
    name = (filename or "").replace("\\", "/").split("/")[-1]
    if "." not in name:
        return ""
    return "." + name.rsplit(".", 1)[-1].lower()


def resolve_vision_prompt(config: Optional[Config], question: str = "") -> str:
    fcfg = _files_cfg(config)
    q = (question or "").strip()
    if q:
        tmpl = (fcfg.vision_focus_prompt or FilesConfig().vision_focus_prompt)
        try:
            return tmpl.format(question=q)
        except (KeyError, ValueError, IndexError):
            return tmpl + "\n" + q
    return fcfg.vision_prompt or FilesConfig().vision_prompt


def _decode_text(data: bytes, max_chars: int) -> Excerpt:
    text = None
    for enc in ("utf-8", "gbk"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        return Excerpt(skipped="text decode failed")
    clipped, truncated = _clip(text, max_chars)
    return Excerpt(text=clipped, truncated=truncated, parser="text")


def _extract_xlsx(data: bytes, max_chars: int) -> Excerpt:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return Excerpt(skipped="openpyxl not installed (pip install sleuth[files])")
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        rows: list[str] = []
        for sheet in wb.worksheets:
            rows.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(cells):
                    rows.append("\t".join(cells))
                if max_chars and sum(len(x) + 1 for x in rows) >= max_chars:
                    break
        wb.close()
    except Exception as exc:
        return Excerpt(skipped=f"xlsx parse failed: {exc}")
    clipped, truncated = _clip("\n".join(rows).strip(), max_chars)
    if not clipped:
        return Excerpt(skipped="xlsx has no extractable text", parser="openpyxl")
    return Excerpt(text=clipped, truncated=truncated, parser="openpyxl")


def _extract_docx(data: bytes, max_chars: int) -> Excerpt:
    try:
        from docx import Document
    except ImportError:
        return Excerpt(skipped="python-docx not installed (pip install sleuth[files])")
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
    except Exception as exc:
        return Excerpt(skipped=f"docx parse failed: {exc}")
    clipped, truncated = _clip("\n".join(parts).strip(), max_chars)
    if not clipped:
        return Excerpt(skipped="docx has no extractable text", parser="python-docx")
    return Excerpt(text=clipped, truncated=truncated, parser="python-docx")


def vision_image_text(
    data: bytes,
    mime: str,
    config: Config,
    *,
    prompt: Optional[str] = None,
) -> str:
    """One-shot vision call. Raises on failure so caller can fall back to OCR."""
    fcfg = _files_cfg(config)
    raw_ref = (fcfg.vision_model or "").strip() or (getattr(config, "model", None) or "")
    if not raw_ref:
        raise RuntimeError("no vision model configured")
    ref = config.prepare_model_ref(raw_ref) if hasattr(config, "prepare_model_ref") else raw_ref
    provider_id, model_id = parse_model_ref(ref)
    from ..provider.factory import build_provider

    provider = build_provider(config, provider_id)
    client = getattr(provider, "_client", None)
    if client is None:
        raise RuntimeError("vision client unavailable")
    mime_use = (mime or "").strip() or "image/jpeg"
    url = f"data:{mime_use};base64,{base64.b64encode(data).decode('ascii')}"
    timeout = max(5.0, float(fcfg.extract_timeout_s or FilesConfig().extract_timeout_s))
    text_prompt = (prompt or "").strip() or resolve_vision_prompt(config)
    resp = client.chat.completions.create(
        model=model_id,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": text_prompt},
                    {"type": "image_url", "image_url": {"url": url}},
                ],
            }
        ],
        timeout=timeout,
    )
    choice = (resp.choices or [None])[0]
    message = getattr(choice, "message", None) if choice is not None else None
    text = (getattr(message, "content", None) or "").strip()
    if not text:
        raise RuntimeError("vision returned empty text")
    return text


def ocr_image_text(data: bytes) -> str:
    engine = None
    try:
        from rapidocr_onnxruntime import RapidOCR

        engine = RapidOCR()
    except ImportError:
        try:
            from rapidocr import RapidOCR

            engine = RapidOCR()
        except ImportError as exc:
            raise RuntimeError("rapidocr not installed (pip install sleuth[ocr])") from exc
    result = engine(data)
    rows = result[0] if isinstance(result, tuple) else result
    if not rows:
        return ""
    texts: list[str] = []
    for row in rows:
        if isinstance(row, dict):
            texts.append(str(row.get("txt") or row.get("text") or ""))
        elif isinstance(row, (list, tuple)) and len(row) >= 2:
            texts.append(str(row[1]))
    return "\n".join(t for t in texts if t).strip()


def _extract_image(
    data: bytes,
    mime: str,
    config: Config,
    max_chars: int,
    *,
    vision_prompt: Optional[str] = None,
) -> Excerpt:
    mode = (_files_cfg(config).image_mode or FilesConfig().image_mode).strip().lower()
    if mode in ("0", "false", "no", "off", "disable", "disabled"):
        return Excerpt(skipped="image parsing disabled")
    prompt = (vision_prompt or "").strip() or resolve_vision_prompt(config)
    errors: list[str] = []
    if mode != "ocr":
        try:
            text = vision_image_text(data, mime, config, prompt=prompt)
            clipped, truncated = _clip(text, max_chars)
            if clipped:
                return Excerpt(text=clipped, truncated=truncated, parser="vision")
            errors.append("vision returned empty text")
        except Exception as exc:
            errors.append(f"vision: {exc}")
    try:
        text = ocr_image_text(data)
        clipped, truncated = _clip(text, max_chars)
        if clipped:
            return Excerpt(text=clipped, truncated=truncated, parser="rapidocr")
        errors.append("rapidocr returned empty text")
    except Exception as exc:
        errors.append(f"rapidocr: {exc}")
    return Excerpt(skipped="image parse failed: " + "; ".join(errors) if errors else "image parse failed")


def render_pdf_pages(data: bytes, config: Optional[Config] = None) -> tuple[list[bytes], str, bool]:
    """Rasterize PDF pages to PNG bytes. Returns (pages, error, more_pages)."""
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return [], "pypdfium2 not installed (pip install sleuth[files])", False
    fcfg = _files_cfg(config)
    try:
        dpi = int(fcfg.pdf_render_dpi or FilesConfig().pdf_render_dpi)
    except (TypeError, ValueError):
        dpi = FilesConfig().pdf_render_dpi
    try:
        max_pages = int(fcfg.pdf_vision_max_pages or FilesConfig().pdf_vision_max_pages)
    except (TypeError, ValueError):
        max_pages = FilesConfig().pdf_vision_max_pages
    dpi = max(36, dpi)
    scale = dpi / 72.0
    try:
        doc = pdfium.PdfDocument(data)
    except Exception as exc:
        return [], f"pdf render failed: {exc}", False
    try:
        total = len(doc)
        limit = min(total, max_pages) if max_pages > 0 else total
        pages: list[bytes] = []
        for i in range(limit):
            page = doc[i]
            bitmap = page.render(scale=scale)
            try:
                pil = bitmap.to_pil()
                buf = io.BytesIO()
                pil.save(buf, format="PNG")
                pages.append(buf.getvalue())
            finally:
                close = getattr(bitmap, "close", None)
                if callable(close):
                    close()
                pclose = getattr(page, "close", None)
                if callable(pclose):
                    pclose()
        return pages, "", total > limit
    except Exception as exc:
        return [], f"pdf render failed: {exc}", False
    finally:
        close = getattr(doc, "close", None)
        if callable(close):
            close()


def _page_prefix(config: Optional[Config], n: int) -> str:
    tmpl = _files_cfg(config).pdf_page_prefix or FilesConfig().pdf_page_prefix
    try:
        return tmpl.format(n=n)
    except (KeyError, ValueError, IndexError):
        return f"# page {n}\n"


def _extract_pdf_images(
    data: bytes,
    config: Config,
    max_chars: int,
    *,
    vision_prompt: Optional[str] = None,
    on_progress=None,
    file_id: str = "",
) -> Excerpt:
    pages, err, more_pages = render_pdf_pages(data, config)
    if err:
        return Excerpt(skipped=err, parser="pypdfium2")
    if not pages:
        return Excerpt(skipped="pdf has no pages to render", parser="pypdfium2")
    parts: list[str] = []
    parsers: list[str] = []
    used = 0
    truncated = bool(more_pages)
    for i, png in enumerate(pages):
        if callable(on_progress):
            fcfg = _files_cfg(config)
            try:
                detail = (fcfg.progress_detail_page or FilesConfig().progress_detail_page).format(
                    page=i + 1, pages=len(pages)
                )
            except (KeyError, ValueError, IndexError):
                detail = f"page {i + 1}/{len(pages)}"
            try:
                on_progress(
                    stage=fcfg.stage_extract_page or "extract_page",
                    file_id=file_id,
                    page=i + 1,
                    pages=len(pages),
                    detail=detail,
                )
            except Exception:
                pass
        remain = max_chars - used if max_chars > 0 else 0
        if max_chars > 0 and remain <= 0:
            truncated = True
            break
        header = _page_prefix(config, i + 1)
        page_ex = _extract_image(
            png,
            "image/png",
            config,
            remain if max_chars > 0 else 0,
            vision_prompt=vision_prompt,
        )
        body = (page_ex.text or "").strip() or (page_ex.skipped or "")
        block = header + body
        if max_chars > 0 and used + len(block) > max_chars:
            take = max_chars - used
            if take <= 0:
                truncated = True
                break
            parts.append(block[:take])
            truncated = True
            if page_ex.parser:
                parsers.append(page_ex.parser)
            break
        parts.append(block)
        used += len(block) + 1
        if page_ex.parser:
            parsers.append(page_ex.parser)
        if page_ex.truncated:
            truncated = True
            break
    text = "\n".join(parts).strip()
    parser = "pypdfium2"
    if parsers:
        uniq = []
        for p in parsers:
            if p not in uniq:
                uniq.append(p)
        parser = "pypdfium2+" + "+".join(uniq)
    if not text:
        return Excerpt(skipped="pdf page vision produced no text", parser=parser)
    return Excerpt(text=text, truncated=truncated, parser=parser)


def _extract_pdf(
    data: bytes,
    max_chars: int,
    config: Config,
    *,
    vision_prompt: Optional[str] = None,
    on_progress=None,
    file_id: str = "",
) -> Excerpt:
    try:
        from pypdf import PdfReader
    except ImportError:
        return Excerpt(skipped="pypdf not installed (pip install sleuth[files])")
    try:
        reader = PdfReader(io.BytesIO(data))
        parts = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:
        return Excerpt(skipped=f"pdf parse failed: {exc}")
    clipped, truncated = _clip("\n".join(parts).strip(), max_chars)
    if clipped:
        return Excerpt(text=clipped, truncated=truncated, parser="pypdf")
    return _extract_pdf_images(
        data,
        config,
        max_chars,
        vision_prompt=vision_prompt,
        on_progress=on_progress,
        file_id=file_id,
    )


def _ext_set(raw: str) -> set:
    return {p.strip().lower() for p in (raw or "").split(",") if p.strip()}


def _kind(data: bytes, mime: str, filename: str, fcfg: FilesConfig) -> str:
    mime_l = (mime or "").strip().lower()
    ext = _ext(filename)
    if mime_l.startswith("image/") or ext in _ext_set(fcfg.image_exts):
        return "image"
    if mime_l == "application/pdf" or ext in _ext_set(fcfg.pdf_exts) or data.startswith(b"%PDF"):
        return "pdf"
    if mime_l in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel.sheet.macroenabled.12",
    ) or ext in _ext_set(fcfg.xlsx_exts):
        return "xlsx"
    if mime_l == "application/vnd.ms-excel" or ext in _ext_set(fcfg.xls_exts):
        return "xls"
    if mime_l in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ) or ext in _ext_set(fcfg.docx_exts):
        return "docx"
    if mime_l.startswith("text/") or mime_l in (
        "application/json",
        "application/xml",
        "application/javascript",
    ) or ext in _ext_set(fcfg.text_exts):
        return "text"
    if data.startswith(b"%PDF"):
        return "pdf"
    if data[:3] == b"\xff\xd8\xff" or data.startswith(b"\x89PNG") or data[8:12] == b"WEBP":
        return "image"
    sniff = 4096
    if b"\x00" not in data[:sniff]:
        return "text"
    return "binary"


def extract_bytes(
    data: bytes,
    *,
    mime: str,
    filename: str,
    config: Optional[Config] = None,
    max_chars: int = 0,
    vision_prompt: Optional[str] = None,
    on_progress=None,
    file_id: str = "",
) -> Excerpt:
    cfg = config or Config()
    fcfg = _files_cfg(cfg)
    limit = int(max_chars or fcfg.excerpt_max_chars or 0)
    if limit <= 0:
        limit = int(FilesConfig().excerpt_max_chars)
    kind = _kind(data, mime, filename, fcfg)
    if kind == "image":
        return _extract_image(
            data, mime, cfg, limit, vision_prompt=vision_prompt
        )
    if kind == "pdf":
        return _extract_pdf(
            data,
            limit,
            cfg,
            vision_prompt=vision_prompt,
            on_progress=on_progress,
            file_id=file_id,
        )
    if kind == "xlsx":
        return _extract_xlsx(data, limit)
    if kind == "xls":
        return Excerpt(skipped="xls not supported; save as xlsx")
    if kind == "docx":
        return _extract_docx(data, limit)
    if kind == "text":
        return _decode_text(data, limit)
    return Excerpt(skipped="unsupported binary type")
